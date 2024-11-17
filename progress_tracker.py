import os
from datetime import datetime
import streamlit as st
import threading
import time
import logging
from PyPDF2 import PdfFileReader as PdfReader
from docx import Document
from bs4 import BeautifulSoup
import markdown
from language_utils import get_translation

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ProgressTracker:
    def __init__(self, total_files):
        self.total_files = total_files
        self.processed_files = 0
        self.current_file = ""
        self.start_time = datetime.now()
        self.embeddings_count = 0
        self.current_chunks = 0
        self.total_chunks = 0
        
        self.file_start_times = {}
        self.file_durations = {}
        
        self.container = st.container()
        with self.container:
            self.progress_bar = st.progress(0)
            self.file_text = st.empty()
            self.status_text = st.empty()
            self.result_text = st.empty()
        
        self.lock = threading.Lock()

    def update(self, file_name, chunks_processed=None, total_chunks=None):
        with self.lock:
            if file_name is not None:
                self.current_file = file_name
                if file_name not in self.file_start_times:
                    self.file_start_times[file_name] = time.time()
                self.file_text.text(f"{get_translation('Document')}: {self.current_file}")
            
            if chunks_processed is not None and total_chunks is not None:
                self.current_chunks = chunks_processed
                self.total_chunks = total_chunks
                progress = chunks_processed / total_chunks if total_chunks > 0 else 0
                self.progress_bar.progress(progress)
                self.status_text.text(get_translation('generating_embeddings').format(
                    processed=chunks_processed,
                    total=total_chunks
                ))

    def increment_file(self):
        with self.lock:
            self.processed_files += 1

    def increment_embeddings(self, count=1):
        with self.lock:
            self.embeddings_count += count

    def get_processing_time(self, file_name):
        with self.lock:
            if file_name in self.file_start_times:
                if file_name in self.file_durations:
                    duration = self.file_durations[file_name]
                else:
                    duration = time.time() - self.file_start_times[file_name]
                return round(duration, 1)
            return 0

    def complete(self):
        with self.lock:
            current_time = time.time()
            for file_name, start_time in self.file_start_times.items():
                if file_name not in self.file_durations:
                    self.file_durations[file_name] = current_time - start_time
            
            self.progress_bar.empty()
            self.file_text.empty()
            self.status_text.empty()

    def get_total_processing_time(self):
        with self.lock:
            total_duration = (datetime.now() - self.start_time).total_seconds()
            return round(total_duration, 1)

    def get_formatted_status(self, file_name):
        with self.lock:
            processing_time = self.get_processing_time(file_name)
            return (f"✨ Processing completed\n\n"
                   f"• Document: {file_name}\n"
                   f"• {self.embeddings_count} embeddings in {processing_time} seconds")

def initialize_tracker(total_files):
    global tracker
    tracker = ProgressTracker(total_files)
    return tracker

def process_document(file_path):
    global tracker
    if tracker:
        tracker.update(os.path.basename(file_path))
        
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    try:
        if file_extension == '.txt':
            result = process_txt(file_path)
            if tracker:
                tracker.increment_file()
            return result
        elif file_extension == '.pdf':
            result = process_pdf(file_path)
            if tracker:
                tracker.increment_file()
            return result
        elif file_extension == '.docx':
            result = process_docx(file_path)
            if tracker:
                tracker.increment_file()
            return result

        elif file_extension == '.html':
            result = process_html(file_path)
            if tracker:
                tracker.increment_file()
            return result
        elif file_extension == '.md':
            result = process_md(file_path)
            if tracker:
                tracker.increment_file()
            return result
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            return None
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {str(e)}", exc_info=True)
        return None

def process_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text = page.extract_text() or ""
            if tracker:
                tracker.increment_embeddings()
            yield text

def process_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
        if tracker:
            tracker.increment_embeddings()
        return text

def process_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    if tracker:
        tracker.increment_embeddings()
    return text

def process_doc(file_path):
    doc = Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    if tracker:
        tracker.increment_embeddings()
    return text

def process_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        soup = BeautifulSoup(file, 'html.parser')
        text = soup.get_text()
        if tracker:
            tracker.increment_embeddings()
        return text

def process_md(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
        html_content = markdown.markdown(content)
        soup = BeautifulSoup(html_content, 'html.parser')
        text = soup.get_text()
        if tracker:
            tracker.increment_embeddings()
        return text

def chunk_text(text_generator, chunk_size=1000, overlap=100):
    current_chunk = ""
    chunk_count = 0
    total_chunks = 0
    
    if hasattr(text_generator, '__iter__') and not isinstance(text_generator, (str, list)):
        text_generator = list(text_generator)
    
    if isinstance(text_generator, (str, list)):
        total_text = "".join(text_generator if isinstance(text_generator, list) else [text_generator])
        total_chunks = max(1, len(total_text) // chunk_size)
        
        if tracker:
            tracker.update(None, chunks_processed=0, total_chunks=total_chunks)

    for text in (text_generator if isinstance(text_generator, list) else [text_generator]):
        current_chunk += text
        while len(current_chunk) >= chunk_size:
            chunk_count += 1
            if tracker:
                tracker.update(None, chunks_processed=chunk_count, total_chunks=total_chunks)
            yield current_chunk[:chunk_size]
            current_chunk = current_chunk[chunk_size-overlap:]
            
    if current_chunk:
        chunk_count += 1
        if tracker:
            tracker.update(None, chunks_processed=chunk_count, total_chunks=total_chunks)
        yield current_chunk

def complete_tracking():
    global tracker
    if tracker:
        tracker.complete()